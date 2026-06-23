from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import StringIO
import re
from deep_translator import GoogleTranslator


textfile_path = 'menu.txt'
template_path = './Templates'
image_path = './Images'
images = {'caution' : '/caution.png'}
output_file_name = 'Menu-Iron_Skillet.docx'
io_folder = "./"

TEXT_WIDTH = 5.65
IMG_WIDTH = 7.5-TEXT_WIDTH

FONT_PROFILES = {
    'name_fr' : {'name' : 'Archivo Black', 'size' : Pt(21), 'bold' : False, 'color' : RGBColor(0, 0, 0)},
    'ingredients_fr' : {'name' : 'Times', 'size' : Pt(12), 'bold' : True, 'color' : RGBColor(0, 0, 0)},
    'name_en' : {'name' : 'Archivo Black', 'size' : Pt(14), 'bold' : False, 'color' : RGBColor(0, 0, 0)},
    'ingredients_en' : {'name' : 'Times', 'size' : Pt(10), 'bold' : True, 'color' : RGBColor(0, 0, 0)},
    'allergens' : {'name' : 'Calibri', 'size' : Pt(10), 'bold' : True, 'color' : RGBColor(255, 0, 0)}
}


TAGS = {
    'vegan' : '/vegan.png',
    'v' : '/vegan.png',
    'vegetarian' : '/vegetarian.png',
    'veg' : '/vegetarian.png',
    'eat well' : '/eat_well.png',
    'ew' : '/eat_well.png',
    'mb' : '/eat_well.png',
    'halal' : '/halal.png',
    'h' : '/halal.png',
    'low carbon' : '/low_carbon.png',
    'lc' : '/low_carbon.png',
    'fe' : '/low_carbon.png'
}


def parse_text(f = None):
    items = []
    if not f:
        f = open(io_folder + textfile_path)
    while True:
        line = f.readline()
        if line == '':
            break
        line = line.strip()
        match = re.fullmatch(r"\d+\. ?(.*)", line)
        if match:
            name = match.group(1).strip()
            ingredients = f.readline().strip()
            if ingredients == '-':
                    ingredients = None
            tags = f.readline().strip().lower()
            if tags == '-':
                    tags = None
            allergens = f.readline().strip()
            if allergens == '-':
                    allergens = None
            items.append({'name':name, 'ingredients':ingredients, 'tags':tags, 'allergens':allergens})
    f.close()
    return items


def set_table_indent(table, inches):
    tblPr = table._element.tblPr
    
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    
    dxa_value = int(inches * 1440)
    tblInd.set(qn('w:w'), str(dxa_value))
    tblInd.set(qn('w:type'), 'dxa')


def apply_font_profile(run, profile_name):
    profile = FONT_PROFILES.get(profile_name)
    if profile:
        run.font.name = profile['name']
        run.font.size = profile['size']
        run.font.bold = profile['bold']
        run.font.color.rgb = profile['color']
        
        
def format_text_paragraphs(text_cell, add_para = True, spacing = 10, space_before = 5, space_after = 5):
    if add_para:
        text_para = text_cell.add_paragraph()
    else:
        text_para = text_cell.paragraphs[0]
    text_para.paragraph_format.line_spacing = Pt(spacing)
    text_para.paragraph_format.space_before = Pt(space_before)
    text_para.paragraph_format.space_after = Pt(space_after)
    return text_para


def create_doc(template_name, items, save=False):
    doc = Document(template_path + template_name)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.9)

    style = doc.styles['normal']

    table = doc.add_table(rows=len(items)+1, cols=2)
    set_table_indent(table, -0.5)

    # Column 1 for the text
    table.columns[0].width = Inches(TEXT_WIDTH)
    # Colume 2 for the label images
    table.columns[1].width = Inches(IMG_WIDTH)

    for i, item in enumerate(items):
        text_cell = table.cell(i+1, 0)
        image_cell = table.cell(i+1, 1)
        print(f"Generating item {i+1}")

        if item['name']:
            p = format_text_paragraphs(text_cell, spacing = 7)
            name_fr = GoogleTranslator(source="en", target="fr").translate(item['name'])
            run = p.add_run(name_fr)
            apply_font_profile(run, 'name_fr')

        if item['ingredients']:
            p = format_text_paragraphs(text_cell, spacing = 8)
            ingredients_fr = GoogleTranslator(source="en", target="fr").translate(item['ingredients'])
            run = p.add_run(ingredients_fr)
            apply_font_profile(run, 'ingredients_fr')

        if itemm['name']:
            p = format_text_paragraphs(text_cell, spacing = 7)
            run = p.add_run(item['name'])
            apply_font_profile(run, 'name_en')

        if item['ingredients']:
            p = format_text_paragraphs(text_cell, spacing = 8, space_after = 0)
            run = p.add_run(item['ingredients'])
            apply_font_profile(run, 'ingredients_en')
        
        if item['allergens']:
            inner_table = text_cell.add_table(rows=1, cols=2)
            
            inner_table.columns[0].width = Inches(0.3)
            inner_table.columns[1].width = Inches(TEXT_WIDTH-0.3)
            
            p = format_text_paragraphs(inner_table.cell(0, 0), add_para = False)
            run = p.add_run()
            run.add_picture(image_path + images['caution'], width=Inches(0.2))
            
            p = format_text_paragraphs(inner_table.cell(0, 1), add_para = False)
            run = p.add_run()
            allergens = item['allergens']
            allergens_fr = GoogleTranslator(source="en", target="fr").translate(allergens)
            run = p.add_run(f'{allergens} / {allergens_fr}')
            apply_font_profile(run, 'allergens')

        if item['tags']:
            add_tags(image_cell, item['tags'], text_cell)

    if save:
        doc.save(io_folder + output_file_name)
    return doc


def add_tags(cell, tags, text_cell):
    tag_list = tags.split(',')
    n_tags = len(tag_list)
    
    rows = 1
    rows = int(n_tags/3 + 1)
    cols = int(n_tags/rows + 0.5)
    table = cell.add_table(rows=rows, cols=cols)
    for i in range(cols):
        table.columns[i].width = Inches(IMG_WIDTH/cols)
    
    for i, tag in enumerate(tag_list[:9]):
        img_path = TAGS.get(tag.strip(), None)
        if img_path:
            try:
                p = format_text_paragraphs(table.cell(int(i/cols), i%cols), add_para = False)
                run = p.add_run()
                run.add_picture(image_path + img_path, width=Inches(0.4))
            except FileNotFoundError:
                print(f"Error: File {img_path} not found")
                
    cell.paragraphs[-1].add_run(' ').font.size = Pt(1)
    text_cell.paragraphs[-1].add_run(' ').font.size = Pt(1)


if __name__ == '__main__':
    items = parse_text()
    create_doc('/is_template.docx', items, save=True)
    print(f"\nSuccessfully generated {output_file_name}\n")
